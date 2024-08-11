import docx
from docx2pdf import convert

import module.google_sheet_helper

doc = docx.Document('IAmAWord.docx')

# run 是換行
# paragraphs 是段落

print("段落數量:", len(doc.paragraphs))
for i in range(google_sheet_helper.get_receiver_info()):
    print(doc.paragraphs[0].add_run(google_sheet_helper.get_receiver_info()[0][0] + " "))
    print(doc.paragraphs[1].add_run(google_sheet_helper.get_receiver_info()[1][0] + " "))

doc.save('IAmAWord.docx')
convert("IAmAWord.docx", "IAmAWord.pdf")